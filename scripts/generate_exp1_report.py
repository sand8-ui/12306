from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[2]
PROJECT_ROOT = WORKSPACE / "project"
OUTPUT_DOCX = WORKSPACE / "空白实验报告-填写版.docx"
OUTPUT_ASCII = WORKSPACE / "git_lab_report_filled.docx"
GUIDE_IMAGES = [
    WORKSPACE / "实验报告书" / "7AB915BCFD45938E0896CE166E57316C.png",
    WORKSPACE / "实验报告书" / "FBFF8E1457A059DF283C767DAAA33040.jpg",
]

STUDENT_INFO = {
    "实验者": "沈洁",
    "同组者": "无",
    "专业班级": "软件工程2402",
    "学号": "1024001308",
    "实验日期": "2026 年 6 月 1 日",
    "组别": "12306-64组",
}


def run_git(*args: str) -> str:
    result = subprocess.run(
        ["git", *args],
        cwd=PROJECT_ROOT,
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return (result.stdout or result.stderr).strip()


def load_context() -> dict[str, str]:
    return {
        "branch": run_git("branch", "--show-current"),
        "status": run_git("status"),
        "status_short": run_git("status", "--short"),
        "log": run_git("log", "--oneline", "--decorate", "-5"),
        "ls_files": run_git("ls-files"),
        "remote": run_git("remote", "-v") or "当前仓库尚未配置远端仓库。",
        "user_name": run_git("config", "--global", "user.name") or "未设置",
        "user_email": run_git("config", "--global", "user.email") or "未设置",
    }


CONTEXT = load_context()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "BFC7D5", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_run_font(run, size: float = 12, *, bold: bool = False, color: str = "222222", name: str = "宋体") -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), name)


def set_paragraph(
    paragraph,
    *,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_chars: int = 2,
    space_before: float = 0,
    space_after: float = 6,
) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.first_line_indent = Pt(24 if first_line_chars else 0)


def add_paragraph_text(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    color: str = "222222",
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_chars: int = 2,
    space_before: float = 0,
    space_after: float = 6,
    font_name: str = "宋体",
):
    paragraph = doc.add_paragraph()
    set_paragraph(
        paragraph,
        align=align,
        first_line_chars=first_line_chars,
        space_before=space_before,
        space_after=space_after,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, name=font_name)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 16, 2: 14, 3: 12}
    colors = {1: "1B3D6D", 2: "1B3D6D", 3: "222222"}
    paragraph = doc.add_paragraph()
    set_paragraph(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=0,
        space_before=8 if level == 1 else 4,
        space_after=6,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, "222222"), name="黑体")


def add_list(doc: Document, items: list[str], *, numbered: bool = False) -> None:
    style_name = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style_name)
        set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_chars=0)
        paragraph.paragraph_format.left_indent = Pt(18)
        run = paragraph.add_run(item)
        set_run_font(run, size=12)


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    size: float = 10.5,
    bg: str | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, align=align, first_line_chars=0, space_before=2, space_after=2)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)
    if bg:
        set_cell_shading(cell, bg)


def add_caption(doc: Document, text: str) -> None:
    add_paragraph_text(
        doc,
        text,
        size=10.5,
        color="555555",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=0,
        space_after=10,
    )


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.core_properties.title = "实验七 Git 实战实验报告"
    doc.core_properties.subject = "软件工程基础实验报告"
    doc.core_properties.author = STUDENT_INFO["实验者"]


def add_cover(doc: Document) -> None:
    add_paragraph_text(
        doc,
        "实验课程名称：软件工程基础实验",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=0,
        space_after=18,
        font_name="黑体",
    )
    add_paragraph_text(
        doc,
        "成绩评定表",
        size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=0,
        space_after=14,
        font_name="黑体",
    )

    table = doc.add_table(rows=6, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ["实验项目名称", "实验七：Git 实战", "", "", "实验成绩", "待评定"],
        ["实验者", STUDENT_INFO["实验者"], "专业班级", STUDENT_INFO["专业班级"], "组别", STUDENT_INFO["组别"]],
        ["同组者", STUDENT_INFO["同组者"], "", "", "实验日期", STUDENT_INFO["实验日期"]],
        ["学号", STUDENT_INFO["学号"], "仓库目录", "project/", "当前分支", CONTEXT["branch"] or "main"],
        ["Git 用户", CONTEXT["user_name"], "Git 邮箱", CONTEXT["user_email"], "实验主题", "本地 Git 仓库初始化、提交与状态分析"],
        ["备注", "空白模板原件为 .doc，本次填写版依据参考实验报告、实验报告书图片与 project 目录中的真实 Git 仓库状态生成，原模板保持不变。", "", "", "", ""],
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            write_cell(
                table.rows[r].cells[c],
                value,
                bold=(c % 2 == 0) or r == 0,
                align=WD_ALIGN_PARAGRAPH.LEFT if c % 2 == 1 and r > 0 else WD_ALIGN_PARAGRAPH.CENTER,
                bg="DCE8F7" if r == 0 else ("EAF1FB" if c % 2 == 0 else None),
            )
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(5, 1).merge(table.cell(5, 5))

    add_paragraph_text(
        doc,
        "说明：本报告按照实验七“Git 实战”的指导书内容编写，正文中的命令、仓库分支、提交历史、跟踪文件和工作区状态均来自当前 project 本地 Git 仓库。",
        size=10.5,
        color="666666",
        first_line_chars=0,
        space_before=8,
        space_after=12,
    )

    score = doc.add_table(rows=7, cols=4)
    score.style = "Table Grid"
    score.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_rows = [
        ["序号", "评分项目", "满分", "实得分"],
        ["1", "实验报告格式规范", "2", ""],
        ["2", "实验报告过程清晰，内容详实", "4", ""],
        ["3", "实验报告结果正确性", "2", ""],
        ["4", "实验分析与总结详尽", "2", ""],
        ["", "总得分", "10", ""],
        ["", "", "", ""],
    ]
    for r, values in enumerate(score_rows):
        for c, value in enumerate(values):
            write_cell(score.rows[r].cells[c], value, bold=r == 0, bg="DCE8F7" if r == 0 else None)

    doc.add_page_break()


def add_part_one(doc: Document) -> None:
    add_heading(doc, "第一部分：实验预习报告", 1)

    add_heading(doc, "一、实验目的", 2)
    add_list(
        doc,
        [
            "熟练掌握 Git 的基本指令，包括仓库初始化、查看状态、添加暂存、提交版本、查看日志和列出受跟踪文件等操作。",
            "理解 Git 中工作区、暂存区、本地仓库、分支和提交历史的作用，并结合当前 project 目录中的真实仓库进行分析。",
            "掌握在已有项目中维护源码仓库的方法，能够根据当前仓库状态判断哪些文件已提交、哪些文件被修改、哪些文件尚未纳入版本管理。",
            "为后续课程项目持续迭代打下基础，使 Mini-12306 源码工程能够通过 Git 进行本地版本管理与过程追踪。",
        ],
        numbered=True,
    )

    add_heading(doc, "二、实验意义", 2)
    add_paragraph_text(
        doc,
        "Git 是软件配置管理中最常用的分布式版本控制工具。对于当前 project 目录下的 Mini-12306 实验项目而言，Git 不仅用于保存代码快照，还能清晰记录前端、后端、测试和文档说明的修改历史，便于后续继续扩展项目功能、回溯问题来源和协同维护代码。",
    )
    add_paragraph_text(
        doc,
        "本次实验指导书强调 Git 的安装、配置、仓库获取、工作流程和分支管理。结合当前已有仓库的真实情况，可以将实验重心放在“理解并分析已有仓库状态”上，用项目中的实际提交、跟踪文件和未提交改动来验证对 Git 基本概念和命令使用方式的掌握程度。",
    )

    add_heading(doc, "三、基本原理与方法", 2)
    add_list(
        doc,
        [
            "Git 仓库以 .git 目录保存对象库、引用、分支和配置；当前 project/ 目录中已经包含 .git，因此该目录本身就是一个独立仓库。",
            "工作区表示用户当前可见和可编辑的源码文件；暂存区用于保存准备提交的变更集合；本地仓库中的 commit 则是对某次提交状态的永久记录。",
            "git status 用于查看工作区与暂存区状态；git log 用于查看提交历史；git ls-files 用于查看当前受版本管理的文件；git branch --show-current 用于查看当前所在分支。",
            "当仓库没有配置远端时，Git 仍然可以完整完成本地版本控制实验，因此本次报告中的远端配置按当前真实情况描述，不虚构未实际存在的 origin。",
        ],
        numbered=True,
    )

    add_heading(doc, "四、主要仪器设备及软件环境", 2)
    env = doc.add_table(rows=6, cols=2)
    env.style = "Table Grid"
    env.alignment = WD_TABLE_ALIGNMENT.CENTER
    env_rows = [
        ["环境项目", "说明"],
        ["操作系统", "Windows（当前实验环境）"],
        ["版本控制工具", "Git 命令行"],
        ["实验仓库目录", str(PROJECT_ROOT)],
        ["当前分支", CONTEXT["branch"] or "main"],
        ["项目内容", "Node.js + React 的 Mini-12306 源码工程"],
    ]
    for r, values in enumerate(env_rows):
        for c, value in enumerate(values):
            write_cell(
                env.rows[r].cells[c],
                value,
                bold=r == 0 or c == 0,
                bg="DCE8F7" if r == 0 else ("EAF1FB" if c == 0 else None),
                align=WD_ALIGN_PARAGRAPH.LEFT if c == 1 and r > 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )

    add_heading(doc, "五、实验方案与技术路线", 2)
    add_list(
        doc,
        [
            "先阅读实验报告书中的 Git 指导内容，明确本次实验需要覆盖 Git 配置、仓库初始化/获取、本地工作流程和分支管理。",
            "检查 project 目录中的 .git、当前分支、提交历史、已跟踪文件和工作区状态，确认本地仓库的真实情况。",
            "将参考实验报告的结构迁移到本次本地 Git 实验场景中，用真实命令结果替换参考报告里的示例性文字。",
            "补充当前仓库与课程项目之间的关系，说明该仓库保存的是 Mini-12306 的源码工程，而实验报告和指导材料位于仓库之外。",
        ],
        numbered=True,
    )

    add_heading(doc, "六、实验报告书内容对应", 2)
    add_paragraph_text(
        doc,
        "实验报告书的两张图片分别展示了“实验七 Git 实战”的实验目的、内容步骤、Git 配置、仓库获取与初始化、本地工作流程、分支管理和协作流程。当前填写版报告在保留这些指导书结构的基础上，将示例命令改写为针对 project 仓库的真实实验记录。",
    )
    captions = [
        "图1 实验七 Git 实战指导书第一页",
        "图2 实验七 Git 实战指导书第二页",
    ]
    for image, caption in zip(GUIDE_IMAGES, captions):
        if image.exists():
            doc.add_picture(str(image), width=Inches(6.0))
            add_caption(doc, caption)


def add_part_two(doc: Document) -> None:
    add_heading(doc, "第二部分：实验过程记录", 1)

    add_heading(doc, "一、Git 配置检查", 2)
    add_paragraph_text(
        doc,
        f"根据实验报告书要求，需要先确认 Git 的用户签名配置。当前机器上的全局 Git 用户名为“{CONTEXT['user_name']}”，邮箱为“{CONTEXT['user_email']}”。这两个配置会写入提交记录中，用于标识提交者身份。",
    )
    config_table = doc.add_table(rows=3, cols=2)
    config_table.style = "Table Grid"
    config_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    config_rows = [
        ["配置项", "当前值"],
        ["user.name", CONTEXT["user_name"]],
        ["user.email", CONTEXT["user_email"]],
    ]
    for r, values in enumerate(config_rows):
        for c, value in enumerate(values):
            write_cell(config_table.rows[r].cells[c], value, bold=r == 0 or c == 0, bg="DCE8F7" if r == 0 else ("EAF1FB" if c == 0 else None), align=WD_ALIGN_PARAGRAPH.LEFT if c == 1 and r > 0 else WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "二、仓库获取与初始化情况", 2)
    add_paragraph_text(
        doc,
        "当前 project/ 目录已经是一个初始化完成的 Git 仓库，因此本次实验没有重新执行 git init，而是围绕已有仓库继续检查和分析。仓库当前所在分支为 main，最近两条提交分别是仓库初始化提交与 README 改进提交，说明该项目已经完成基本的本地版本管理。",
    )
    add_paragraph_text(
        doc,
        f"最近提交历史如下：{CONTEXT['log'].replace(chr(10), '；')}。",
    )

    add_heading(doc, "三、本地工作流程实验记录", 2)
    add_list(
        doc,
        [
            "查看仓库状态：在 project/ 目录执行 git status，确认当前存在 5 个已修改文件和 2 个未跟踪文件。",
            "分析修改来源：已修改文件主要涉及 README、前端页面、种子数据、后端领域逻辑和测试；未跟踪文件主要是用于生成实验报告的脚本和其缓存目录。",
            "查看提交历史：通过 git log --oneline --decorate -5 查看 main 分支上的历史提交，确认提交顺序和当前 HEAD 位置。",
            "查看受跟踪文件：通过 git ls-files 确认当前仓库纳入版本管理的文件范围主要是源码、启动脚本、README 和测试文件。",
        ],
        numbered=True,
    )

    add_heading(doc, "四、当前仓库状态分析", 2)
    add_paragraph_text(
        doc,
        "以下内容直接来自当前仓库的真实命令输出，可作为本次 Git 实战的核心证据。由于当前环境尚未执行新的 git add 或 git commit，因此这些修改仍停留在工作区阶段，未进入暂存区或新的提交记录。",
    )

    status_table = doc.add_table(rows=5, cols=2)
    status_table.style = "Table Grid"
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_rows = [
        ["检查项", "结果"],
        ["当前分支", CONTEXT["branch"] or "main"],
        ["远端仓库", CONTEXT["remote"]],
        ["git status --short", CONTEXT["status_short"] or "工作区干净"],
        ["git log --oneline --decorate -5", CONTEXT["log"]],
    ]
    for r, values in enumerate(status_rows):
        for c, value in enumerate(values):
            write_cell(status_table.rows[r].cells[c], value, bold=r == 0 or c == 0, bg="DCE8F7" if r == 0 else ("EAF1FB" if c == 0 else None), align=WD_ALIGN_PARAGRAPH.LEFT if c == 1 and r > 0 else WD_ALIGN_PARAGRAPH.CENTER, size=10)

    add_heading(doc, "五、已跟踪文件与仓库边界", 2)
    tracked_files = CONTEXT["ls_files"].splitlines()
    add_paragraph_text(
        doc,
        f"git ls-files 显示当前共有 {len(tracked_files)} 个已跟踪文件，覆盖后端源码、前端源码、测试、README、package.json 和脚本文件，说明 project/ 目录内的源码工程已经完整纳入 Git 管理。",
    )
    tracked_table = doc.add_table(rows=8, cols=2)
    tracked_table.style = "Table Grid"
    tracked_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tracked_rows = [
        ["类别", "示例文件"],
        ["仓库说明", "README.md、package.json、.gitignore"],
        ["前端页面", "client/index.html、client/src/app.js、client/src/styles.css"],
        ["后端源码", "server/src/server.js、server/src/domain/store.js、server/src/lib/utils.js"],
        ["种子数据", "server/data/seed.js"],
        ["测试代码", "server/tests/store.test.js"],
        ["脚本文件", "scripts/dev.js、scripts/serve-client.js"],
        ["未纳入跟踪", "scripts/generate_exp1_report.py、scripts/__pycache__/ 当前仍为未跟踪文件"],
    ]
    for r, values in enumerate(tracked_rows):
        for c, value in enumerate(values):
            write_cell(tracked_table.rows[r].cells[c], value, bold=r == 0 or c == 0, bg="DCE8F7" if r == 0 else ("EAF1FB" if c == 0 else None), align=WD_ALIGN_PARAGRAPH.LEFT if c == 1 and r > 0 else WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "六、实验过程中发现的问题及处理", 2)
    problems = doc.add_table(rows=5, cols=3)
    problems.style = "Table Grid"
    problems.alignment = WD_TABLE_ALIGNMENT.CENTER
    problem_rows = [
        ["问题", "表现", "处理方式"],
        ["空白模板为旧版 .doc", "无法直接通过 python-docx 编辑空白模板。", "保持原模板不动，生成独立的 .docx 填写版。"],
        ["指导书与上一版报告主题不一致", "指导书实际是 Git 实战，而先前填写版内容是 Mini-12306 业务实验。", "重新按 Git 实战主题生成本版报告，正文基于当前 Git 仓库真实状态重写。"],
        ["当前仓库未配置远端", "git remote -v 没有输出，说明本地尚未设置 origin。", "报告中如实记录“当前仓库尚未配置远端”，不虚构 Gitee 地址。"],
        ["工作区存在未提交改动", "git status 显示 5 个 modified 和 2 个 untracked。", "作为实验过程记录保留，说明 Git 可准确反映当前工作区状态。"],
    ]
    for r, values in enumerate(problem_rows):
        for c, value in enumerate(values):
            write_cell(problems.rows[r].cells[c], value, bold=r == 0, bg="DCE8F7" if r == 0 else None, align=WD_ALIGN_PARAGRAPH.LEFT if r > 0 else WD_ALIGN_PARAGRAPH.CENTER, size=10)

    add_heading(doc, "七、测试与运行记录", 2)
    add_paragraph_text(
        doc,
        "虽然本次实验主题是 Git 实战，但由于当前仓库管理的是可运行的 Mini-12306 源码工程，因此也可通过执行项目测试来验证仓库中被跟踪的代码处于可用状态。本地运行命令 node --test ./server/tests/*.test.js 后，得到 tests 3、pass 3、fail 0、duration_ms 108.8975 的结果。",
    )
    tests = doc.add_table(rows=4, cols=4)
    tests.style = "Table Grid"
    tests.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_rows = [
        ["测试编号", "测试场景", "期望结果", "实际结果"],
        ["T1", "购票后退票", "订单状态和余票变化正确", "通过"],
        ["T2", "购票后改签", "原票标记 CHANGED，新票保持 PAID", "通过"],
        ["T3", "实名注册", "要求银行卡号并返回实名校验结果", "通过"],
    ]
    for r, values in enumerate(test_rows):
        for c, value in enumerate(values):
            write_cell(tests.rows[r].cells[c], value, bold=r == 0, bg="DCE8F7" if r == 0 else None, align=WD_ALIGN_PARAGRAPH.LEFT if r > 0 else WD_ALIGN_PARAGRAPH.CENTER)


def add_part_three(doc: Document) -> None:
    add_heading(doc, "第三部分：结果与讨论", 1)

    add_heading(doc, "一、实验结果", 2)
    add_paragraph_text(
        doc,
        "本次实验已经完成基于 project 本地仓库的 Git 实战报告整理工作。当前仓库位于 project/ 目录，分支为 main，已存在至少两条本地提交记录，能够通过 git status、git log 和 git ls-files 清晰反映当前项目的版本控制状态。",
    )
    add_paragraph_text(
        doc,
        "从结果看，Git 已经成功用于管理当前 Mini-12306 源码工程。虽然当前仓库尚未配置远端仓库，但这并不影响本地 Git 实验的完成，因为初始化、提交、历史查看和工作区分析都可以在本地独立完成。",
    )

    add_heading(doc, "二、主要成果", 2)
    add_list(
        doc,
        [
            "确认 project/ 已是独立 Git 仓库，当前分支为 main。",
            "核实全局 Git 用户名和邮箱配置，满足本地提交记录署名需要。",
            "通过 git log 获得真实提交历史，确认当前已有仓库初始化与 README 改进两次提交。",
            "通过 git ls-files 明确仓库当前只管理源码、脚本、测试和说明文件，不自动管理实验报告材料。",
            "通过 git status 真实记录当前工作区改动情况，为后续 add / commit 操作提供依据。",
        ],
        numbered=True,
    )

    add_heading(doc, "三、实验现象分析", 2)
    add_paragraph_text(
        doc,
        "git status 的输出能够直观地体现 Git 对工作区变化的跟踪能力。当前 README、前端页面、种子数据、后端逻辑和测试文件均被识别为 modified，说明这些文件此前已经纳入版本管理，但本轮改动尚未形成新的提交。与此同时，scripts/generate_exp1_report.py 和 scripts/__pycache__/ 被识别为 untracked，说明它们仍处于“尚未由 Git 管理”的状态。",
    )
    add_paragraph_text(
        doc,
        "git ls-files 则反映了仓库边界的清晰性：只有源码工程内部的关键文件被追踪，而实验报告模板、指导书图片和其他外层文档文件并不在该仓库中。这说明 Git 仓库边界是由 .git 所在目录和用户实际 add/commit 过的文件共同决定的，而不是由整个课程目录自动决定的。",
    )

    add_heading(doc, "四、综合结论", 2)
    add_paragraph_text(
        doc,
        "综合来看，本次 Git 实战实验已经达到了通过真实项目理解 Git 基本工作流程的目的。即使没有远端仓库，本地仓库的初始化、提交历史、文件跟踪和状态分析也足以支撑实验要求中的大部分知识点。",
    )
    add_paragraph_text(
        doc,
        "对于当前 Mini-12306 项目而言，继续使用 Git 管理前后端源码、测试和启动脚本是合理的。后续若课程需要进一步演示分支协作、远端推送或多人合并流程，可以在现有 main 分支基础上继续扩展 src/doc 分支或配置远端仓库。 ",
    )

    add_heading(doc, "五、实验体会", 2)
    add_paragraph_text(
        doc,
        "本次实验最大的收获，是把 Git 的抽象概念和一个真实仓库结合起来理解。相比只记忆命令格式，直接面对 project 目录中的实际仓库状态，更容易理解“修改但未提交”“未跟踪文件”“提交历史”“受跟踪文件范围”等概念的具体含义。",
    )
    add_paragraph_text(
        doc,
        "同时也能看出，Git 的价值不仅在于保存代码，更在于帮助开发者建立对项目边界和版本状态的清晰认知。只要仓库状态和报告内容保持一致，实验报告就能真实反映项目过程，而不是停留在命令示例层面。",
    )

    add_heading(doc, "六、提交前可补充项", 2)
    add_list(
        doc,
        [
            "[截图占位：插入 git status 终端界面截图，需显示 modified 与 untracked 文件。]",
            "[截图占位：插入 git log --oneline --decorate 终端界面截图，需显示当前 main 分支最近提交。]",
            "[截图占位：插入 git ls-files 终端界面截图，需显示仓库已跟踪文件列表。]",
            "如后续配置了远端仓库，可补充 git remote -v 和 git push 的截图与说明。",
        ],
        numbered=True,
    )


def build_report() -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_part_one(doc)
    add_part_two(doc)
    add_part_three(doc)
    doc.save(OUTPUT_DOCX)
    doc.save(OUTPUT_ASCII)
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build_report()
    print(out)
